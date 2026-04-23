"""
Convert paper_draft_zh.md -> paper_draft_zh.docx
with proper heading styles, tables, and math placeholder handling.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
if len(sys.argv) >= 3:
    SRC = Path(sys.argv[1])
    DST = Path(sys.argv[2])
else:
    SRC = Path("/home/xut/csclip/paper_draft_zh.md")
    DST = Path("/home/xut/csclip/paper_draft_zh.docx")


def set_font(run, name="Times New Roman", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt([0, 16, 14, 13, 12][min(level, 4)])
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), "黑体")
        rPr.insert(0, rFonts)
    return p


def add_paragraph(doc, text, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run)
    return p


def parse_table(lines):
    """Parse markdown table lines into list of row lists."""
    rows = []
    for line in lines:
        if re.match(r"\|[-: |]+\|", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= ncols:
                break
            cell = table.cell(i, j)
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.4)


def main():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            add_heading(doc, line[2:].strip(), 1)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
            i += 1
            continue

        # H4
        if line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 4)
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        # Table
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            add_table(doc, rows)
            continue

        # Horizontal rule / front matter separator
        if line.strip() in ("---", "***", "==="):
            doc.add_paragraph()
            i += 1
            continue

        # Bold **...** metadata lines at top
        # List items
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            content = line.strip()[2:]
            # Handle bold inside list
            parts = re.split(r"(\*\*.*?\*\*)", content)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    set_font(run, bold=True)
                else:
                    run = p.add_run(part)
                    set_font(run)
            i += 1
            continue

        # Math lines ($$...$$) - render as italic placeholder
        if line.strip().startswith("$$"):
            math_lines = [line]
            if not line.strip().endswith("$$") or line.strip() == "$$":
                i += 1
                while i < len(lines) and not lines[i].strip().endswith("$$"):
                    math_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    math_lines.append(lines[i])
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" ".join(l.strip() for l in math_lines))
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.name = "Cambria Math"
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph - handle inline formatting
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # split by bold
        parts = re.split(r"(\*\*.*?\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                set_font(run, bold=True)
            else:
                # handle inline code `...`
                sub_parts = re.split(r"(`[^`]+`)", part)
                for sp in sub_parts:
                    if sp.startswith("`") and sp.endswith("`"):
                        run = p.add_run(sp[1:-1])
                        run.font.name = "Courier New"
                        run.font.size = Pt(10)
                    else:
                        run = p.add_run(sp)
                        set_font(run)
        i += 1

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()

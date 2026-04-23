"""
Convert patent_draft_zh.md -> patent_draft_zh.docx
专门为专利文档优化：
  1. 公式（含 Σ μ ‖ ← × ² ¹ ∈ ⟨ ⟩ 等 Unicode 数学符号的行）居中、Cambria Math 字体渲染
  2. "图 N 为..." 的附图说明后自动插入图片占位符方框
  3. 保留原有标题、表格、列表、加粗、代码块处理
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if len(sys.argv) >= 3:
    SRC = Path(sys.argv[1])
    DST = Path(sys.argv[2])
else:
    SRC = Path("/home/xut/csclip/patent_draft_zh.md")
    DST = Path("/home/xut/csclip/patent_draft_zh.docx")

# ---------------- 字体工具 ---------------- #

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False,
             color=None, east_asia="宋体"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), east_asia)
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = [0, 18, 15, 13, 12, 12]
    for run in p.runs:
        set_font(run, name="Times New Roman", size=sizes[min(level, 5)],
                 bold=True, east_asia="黑体")
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

# ---------------- 公式判定 ---------------- #

MATH_HINT_CHARS = set("ΣμρπγτεαβΩΔ∇∑∏∈∪∩⟨⟩‖←→×÷²¹⁰·∘ŷ𝟙")


CHINESE_PUNCT = set("，。；：！？、""''（）《》【】")


def looks_like_formula(line: str) -> bool:
    """判断一行是否为纯数学公式（排除含中文标点或以中文字开头的句子）。"""
    s = line.strip()
    if not s or s.startswith(("#", "|", "-", "*", ">", "```")):
        return False
    if s.startswith("**"):
        return False
    # 排除含中文标点（中文语句常带）
    if any(ch in CHINESE_PUNCT for ch in s):
        return False
    # 排除以中文字符开头的行
    if re.match(r"^[\u4e00-\u9fff]", s):
        return False
    # 含 Unicode 数学字符
    if any(ch in MATH_HINT_CHARS for ch in s):
        return True
    # 形如 R_feat(i, j) = ... 或 f = 0.9 × ...
    if re.match(r"^[A-Za-zΣμρπγτε_^]+(\([^)]*\))?(\^\S+)?(_\{[^}]+\})?\s*[=←]", s):
        return True
    return False


def add_formula(doc, text: str, number: str | None = None):
    """居中插入公式行，可选公式编号。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.strip())
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    run.font.italic = True
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rPr.insert(0, rFonts)
    if number:
        tab = p.add_run(f"\t({number})")
        tab.font.name = "Times New Roman"
        tab.font.size = Pt(11)
    return p

# ---------------- 图片占位符 ---------------- #

FIG_PATTERN = re.compile(r"^图\s*([0-9]+)\s*为(.+?)[；。]?$")


def add_figure_placeholder(doc, fig_no: str, caption: str):
    """为图 N 插入一个占位表格 + 图题。"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置单元格高度
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), "2400")  # ~4cm
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(f"[ 此处预留：图 {fig_no} 占位，插入后请删除本行 ]")
    set_font(r, name="Times New Roman", size=10, italic=True,
             color=(128, 128, 128))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"图 {fig_no}  {caption.strip()}")
    set_font(run, name="Times New Roman", size=10.5, bold=True)
    doc.add_paragraph()  # spacing


# ---------------- 行内格式 ---------------- #

def emit_inline(p, text: str):
    """处理加粗 **x** 和行内代码 `x`。"""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = p.add_run(part[2:-2])
            set_font(run, bold=True)
        else:
            sub = re.split(r"(`[^`]+`)", part)
            for sp in sub:
                if sp.startswith("`") and sp.endswith("`"):
                    run = p.add_run(sp[1:-1])
                    run.font.name = "Courier New"
                    run.font.size = Pt(10.5)
                else:
                    run = p.add_run(sp)
                    set_font(run)


def add_paragraph(doc, text: str, indent: bool = True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 2 个汉字
    emit_inline(p, text)
    return p

# ---------------- 表格 ---------------- #

def parse_table(lines):
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-: |]+\|\s*$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 去掉 markdown 粗体符号
            bold = False
            txt = cell_text
            if txt.startswith("**") and txt.endswith("**"):
                txt = txt[2:-2]
                bold = True
            run = para.add_run(txt)
            set_font(run, name="Times New Roman", size=10.5,
                     bold=(i == 0 or bold))


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)

# ---------------- 主流程 ---------------- #

def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # 正文默认字体
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        # 代码块
        if stripped.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            add_code_block(doc, code)
            i += 1
            continue

        # 表格
        if stripped.startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            add_table(doc, parse_table(tbl))
            continue

        # 水平分隔线
        if stripped in ("---", "***", "==="):
            doc.add_paragraph()
            i += 1
            continue

        # 列表
        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            emit_inline(p, stripped[2:])
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 图 N 为 xxx；- 在"附图说明"段落后自动插入占位
        fig = FIG_PATTERN.match(stripped)
        if fig:
            add_figure_placeholder(doc, fig.group(1), fig.group(2))
            i += 1
            continue

        # 公式行
        if looks_like_formula(stripped):
            # 合并连续的公式行
            block = [stripped]
            i += 1
            while i < len(lines) and looks_like_formula(lines[i].strip()):
                block.append(lines[i].strip())
                i += 1
            for f in block:
                add_formula(doc, f)
            continue

        # 普通段落
        add_paragraph(doc, line.rstrip(), indent=True)
        i += 1

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()

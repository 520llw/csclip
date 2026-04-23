#!/usr/bin/env python3
"""
将 paper_full_zh.md 转换为 Word (.docx) 格式
适用于 CMIG 投稿前的格式检查
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


def set_chinese_font(run, font_name='SimSun', font_size=10.5):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_english_font(run, font_name='Times New Roman', font_size=10.5):
    """设置英文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)


def add_formula_placeholder(doc, formula_text):
    """添加公式占位符（Word中公式需要手动或MathType插入）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[公式: {formula_text}]")
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    set_chinese_font(run, 'SimSun', 10)


def process_markdown_to_docx(md_file, output_file):
    doc = Document()
    
    # 设置页面大小（A4）
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 简单解析：按行处理
    lines = content.split('\n')
    in_code_block = False
    in_table = False
    table_lines = []
    
    for line in lines:
        stripped = line.strip()
        
        # 代码块
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            if not in_code_block:
                # 结束代码块，添加内容
                if table_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    run = p.add_run('\n'.join(table_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    table_lines = []
            continue
        
        if in_code_block:
            table_lines.append(line)
            continue
        
        # 跳过 YAML front matter
        if stripped == '---':
            continue
        
        # 标题
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 16)
                run.bold = True
        elif stripped.startswith('## '):
            text = stripped[3:]
            # 跳过占位符说明
            if '[图占位]' in text or '[数据占位]' in text:
                continue
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 14)
                run.bold = True
        elif stripped.startswith('### '):
            text = stripped[4:]
            if '[图占位]' in text or '[数据占位]' in text:
                continue
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 12)
                run.bold = True
        elif stripped.startswith('#### '):
            text = stripped[5:]
            if '[图占位]' in text or '[数据占位]' in text:
                continue
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                set_chinese_font(run, 'SimHei', 11)
                run.bold = True
        
        # 表格（简化处理：检测 | 分隔）
        elif '|' in stripped and not stripped.startswith('['):
            # 这是表格行
            if not in_table:
                in_table = True
                table_lines = [stripped]
            else:
                table_lines.append(stripped)
        
        elif in_table and not stripped.startswith('|'):
            # 表格结束，创建表格
            if len(table_lines) >= 2:
                create_table_from_lines(doc, table_lines)
            in_table = False
            table_lines = []
            # 继续处理当前行
            if stripped:
                add_paragraph_with_format(doc, line)
        
        # 公式块 ($$ ... $$)
        elif stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
            formula = stripped[2:-2].strip()
            add_formula_placeholder(doc, formula[:50] + '...' if len(formula) > 50 else formula)
        
        # 普通段落
        elif stripped:
            if '[图占位]' in stripped or '[数据占位]' in stripped:
                # 添加占位符提示
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.italic = True
                set_chinese_font(run, 'SimSun', 10)
            else:
                add_paragraph_with_format(doc, line)
    
    # 处理最后的表格
    if in_table and len(table_lines) >= 2:
        create_table_from_lines(doc, table_lines)
    
    doc.save(output_file)
    print(f"Word 文档已保存至: {output_file}")


def create_table_from_lines(doc, lines):
    """从 markdown 表格行创建 Word 表格"""
    # 解析表头、分隔行、数据行
    header_line = lines[0]
    # 跳过分隔行（包含 --- 的行）
    data_lines = [l for l in lines[1:] if '---' not in l and l.strip()]
    
    if not data_lines:
        return
    
    # 解析列
    headers = [c.strip() for c in header_line.split('|') if c.strip()]
    num_cols = len(headers)
    num_rows = len(data_lines) + 1  # +1 for header
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # 填充表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                set_chinese_font(r, 'SimHei', 10)
    
    # 填充数据
    for i, data_line in enumerate(data_lines):
        cells = [c.strip() for c in data_line.split('|') if c.strip()]
        for j, cell_text in enumerate(cells[:num_cols]):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                for r in p.runs:
                    set_chinese_font(r, 'SimSun', 10)


def add_paragraph_with_format(doc, line):
    """添加带格式的段落，处理行内格式"""
    # 检测粗体 **text**
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|\$[^$]+\$)', line)
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    p.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    
    for part in parts:
        if not part:
            continue
        run = p.add_run()
        
        if part.startswith('**') and part.endswith('**'):
            # 粗体
            text = part[2:-2]
            run.text = text
            run.bold = True
            set_chinese_font(run, 'SimHei', 10.5)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # 斜体
            text = part[1:-1]
            run.text = text
            run.italic = True
            set_chinese_font(run, 'SimSun', 10.5)
        elif part.startswith('$') and part.endswith('$'):
            # 行内公式
            run.text = part
            run.font.name = 'Times New Roman'
            run.font.italic = True
        else:
            # 普通文本
            run.text = part
            # 检测是否主要是英文
            if len(re.findall(r'[a-zA-Z]', part)) > len(part) * 0.5:
                set_english_font(run, 'Times New Roman', 10.5)
            else:
                set_chinese_font(run, 'SimSun', 10.5)


if __name__ == '__main__':
    process_markdown_to_docx('/home/xut/csclip/paper_full_zh.md', 
                              '/home/xut/csclip/BALF_Analyzer_Paper_v1.docx')
